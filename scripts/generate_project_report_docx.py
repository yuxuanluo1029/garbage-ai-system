from __future__ import annotations

import csv
from pathlib import Path

import yaml
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
BASE_DIR = PROJECT_ROOT.parent
EXPERIMENT_DIR = BASE_DIR / "实验图片"
OUTPUT_DOCX = BASE_DIR / "垃圾分类智能站项目说明书.docx"
TRAINING_CURVE = EXPERIMENT_DIR / "training_metrics_curve.png"


def set_east_asia_font(style_or_run, font_name: str = "宋体", size: float | None = None, bold: bool | None = None) -> None:
    style_or_run.font.name = font_name
    if size is not None:
        style_or_run.font.size = Pt(size)
    if bold is not None:
        style_or_run.font.bold = bold
    style_or_run.font.color.rgb = RGBColor(0, 0, 0)
    element = style_or_run._element
    rpr = element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.start_type = WD_SECTION.NEW_PAGE

    normal = document.styles["Normal"]
    set_east_asia_font(normal, size=10.5, bold=False)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Cm(0.74)

    title = document.styles["Title"]
    set_east_asia_font(title, size=18, bold=True)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading1 = document.styles["Heading 1"]
    set_east_asia_font(heading1, size=14, bold=True)
    heading1.paragraph_format.space_before = Pt(12)
    heading1.paragraph_format.space_after = Pt(6)
    heading1.paragraph_format.first_line_indent = Cm(0)

    heading2 = document.styles["Heading 2"]
    set_east_asia_font(heading2, size=12, bold=True)
    heading2.paragraph_format.space_before = Pt(10)
    heading2.paragraph_format.space_after = Pt(4)
    heading2.paragraph_format.first_line_indent = Cm(0)


def add_title(document: Document, title_text: str, subtitle: str) -> None:
    paragraph = document.add_paragraph(style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title_text)
    set_east_asia_font(run, size=18, bold=True)

    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(subtitle)
    set_east_asia_font(run, size=10.5, bold=False)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def add_heading(document: Document, text: str, level: int) -> None:
    style = "Heading 1" if level == 1 else "Heading 2"
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_east_asia_font(run, size=14 if level == 1 else 12, bold=True)
    paragraph.paragraph_format.first_line_indent = Cm(0)


def add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Normal")
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    run = paragraph.add_run(text)
    set_east_asia_font(run, size=10.5, bold=False)


def add_caption(document: Document, text: str, prefix: str = "图") -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run(f"{prefix} {text}")
    set_east_asia_font(run, size=10.5, bold=False)


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("left", "top", "right", "bottom"):
        edge_data = kwargs.get(edge, {"val": "nil"})
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def style_triple_line_table(table) -> None:
    rows = table.rows
    for row in rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                left={"val": "nil"},
                right={"val": "nil"},
                top={"val": "nil"},
                bottom={"val": "nil"},
            )

    for cell in rows[0].cells:
        set_cell_border(
            cell,
            left={"val": "nil"},
            right={"val": "nil"},
            top={"val": "single", "sz": 10, "color": "000000"},
            bottom={"val": "single", "sz": 6, "color": "000000"},
        )
    for cell in rows[-1].cells:
        set_cell_border(
            cell,
            left={"val": "nil"},
            right={"val": "nil"},
            top={"val": "nil"},
            bottom={"val": "single", "sz": 10, "color": "000000"},
        )


def fill_cell_text(cell, text: str, bold: bool = False, align: WD_PARAGRAPH_ALIGNMENT = WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run(str(text))
    set_east_asia_font(run, size=10.5, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_triple_line_table(document: Document, caption: str, headers: list[str], rows: list[list[str]], widths_cm: list[float]) -> None:
    add_caption(document, caption, prefix="表")
    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for idx, width in enumerate(widths_cm):
        for row in table.rows:
            row.cells[idx].width = Cm(width)

    for idx, header in enumerate(headers):
        fill_cell_text(table.rows[0].cells[idx], header, bold=True)

    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if col_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            fill_cell_text(table.rows[row_idx].cells[col_idx], value, bold=False, align=align)

    style_triple_line_table(table)
    document.add_paragraph()


def add_picture(document: Document, path: Path, width_cm: float, caption: str) -> None:
    if not path.exists():
        return
    document.add_picture(str(path), width=Cm(width_cm))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(document, caption, prefix="图")


def create_training_curve(results_csv: Path, output_path: Path) -> Path | None:
    try:
        import matplotlib.pyplot as plt
    except Exception:
        return None

    epochs, precision, recall, map50, map5095 = [], [], [], [], []
    with results_csv.open(encoding="utf-8") as file:
        reader = csv.DictReader(file)
        for row in reader:
            epochs.append(int(float(row["epoch"])))
            precision.append(float(row["metrics/precision(B)"]))
            recall.append(float(row["metrics/recall(B)"]))
            map50.append(float(row["metrics/mAP50(B)"]))
            map5095.append(float(row["metrics/mAP50-95(B)"]))

    plt.figure(figsize=(10, 4.8))
    plt.plot(epochs, map50, label="mAP50", linewidth=2.2, color="#1f7a4d")
    plt.plot(epochs, map5095, label="mAP50-95", linewidth=2.2, color="#2b6cb0")
    plt.plot(epochs, precision, label="Precision", linewidth=1.8, color="#d97706")
    plt.plot(epochs, recall, label="Recall", linewidth=1.8, color="#7c3aed")
    plt.xlabel("Epoch")
    plt.ylabel("Metric Value")
    plt.title("YOLOv8 Training Metrics")
    plt.grid(alpha=0.25, linestyle="--")
    plt.legend()
    plt.tight_layout()
    plt.savefig(output_path, dpi=220)
    plt.close()
    return output_path


def read_dataset_info() -> dict:
    dataset_dir = EXPERIMENT_DIR / "dataset"
    yaml_path = dataset_dir / "dataset.yaml"
    args_path = EXPERIMENT_DIR / "args.yaml"
    results_csv = EXPERIMENT_DIR / "results.csv"

    with yaml_path.open(encoding="utf-8") as file:
        dataset_cfg = yaml.safe_load(file)
    with args_path.open(encoding="utf-8") as file:
        args = yaml.safe_load(file)

    train_count = sum(1 for _ in (dataset_dir / "train.txt").open(encoding="utf-8"))
    val_count = sum(1 for _ in (dataset_dir / "val.txt").open(encoding="utf-8"))
    test_count = sum(1 for _ in (dataset_dir / "test.txt").open(encoding="utf-8"))

    rows = list(csv.DictReader(results_csv.open(encoding="utf-8")))
    best_map50 = max(rows, key=lambda row: float(row["metrics/mAP50(B)"]))
    best_map95 = max(rows, key=lambda row: float(row["metrics/mAP50-95(B)"]))
    best_precision = max(rows, key=lambda row: float(row["metrics/precision(B)"]))
    best_recall = max(rows, key=lambda row: float(row["metrics/recall(B)"]))
    final_row = rows[-1]

    return {
        "dataset_cfg": dataset_cfg,
        "args": args,
        "train_count": train_count,
        "val_count": val_count,
        "test_count": test_count,
        "total_count": train_count + val_count + test_count,
        "epochs_run": len(rows),
        "best_map50": best_map50,
        "best_map95": best_map95,
        "best_precision": best_precision,
        "best_recall": best_recall,
        "final_row": final_row,
    }


def build_report() -> Path:
    info = read_dataset_info()
    document = Document()
    configure_document(document)

    add_title(
        document,
        "垃圾分类智能站项目说明书",
        "人工智能导论课程大作业",
    )

    add_triple_line_table(
        document,
        "1 项目基本信息",
        ["项目项", "内容"],
        [
            ["项目名称", "垃圾分类智能站：基于 YOLOv8 与通义千问的校园垃圾识别与智能分析系统"],
            ["核心算法", "YOLOv8s 目标检测模型 + 阿里云通义千问 API 智能体解释"],
            ["系统形态", "登录注册入口 + 垃圾识别栏目 + 个性推荐栏目 + 分类博客栏目 + 个人信息栏目"],
            ["数据规模", f"{info['total_count']} 张图像，四类垃圾检测任务"],
            ["部署方式", "本地 Web 演示可运行，并提供 GitHub + Render 公网部署方案"],
            ["文档说明", "成员信息栏位先保留为待补充，便于后续按实际小组情况填写"],
        ],
        [3.5, 12.5],
    )

    add_heading(document, "一、问题介绍", level=1)
    add_body_paragraph(
        document,
        "本项目面向校园日常垃圾分类场景。实际生活中，宿舍楼下、教学楼走廊和食堂附近的垃圾桶虽然普遍张贴了分类标识，但很多同学在投放时仍然会遇到两个典型问题：一是面对果皮、纸箱、塑料盒、废电池、充电宝这类边界较强的物品时，往往凭经验判断，容易出错；二是即使知道大类，也不清楚一些特殊物品是否需要单独送到指定回收点。"
    )
    add_body_paragraph(
        document,
        "因此，本项目希望把“看图识别 + 智能解释 + 学习推荐 + 交流记录”结合起来，做成一个更接近真实使用场景的网页系统。用户在系统中可以先上传垃圾图片获取四类垃圾识别结果，再通过智能体助手追问原因；如果想继续学习，还可以进入个性推荐栏目观看垃圾分类相关视频，或者在博客栏目中记录和分享自己的分类实践。"
    )
    add_body_paragraph(
        document,
        "从课程角度看，这个选题的优点是问题明确、生活相关性强、展示形式直观，既体现了深度学习目标检测算法的使用，也体现了生成式大模型在解释和问答方面的补充价值，适合在 7 分钟 PPT 中完整呈现项目过程、技术细节和结果意义。"
    )
    add_triple_line_table(
        document,
        "2 项目问题与对应解决思路",
        ["实际问题", "对应解决思路"],
        [
            ["常见垃圾容易混投", "使用 YOLOv8 对图片中的垃圾目标进行检测，并直接输出四类垃圾判断结果"],
            ["特殊物品规则难记", "接入阿里云通义千问 API，让智能体对识别结果进行二次解释和补充说明"],
            ["课堂展示只讲模型太单薄", "把识别、推荐、博客、个人信息整合成栏目化网页，增强可视化效果"],
            ["模型输出不一定覆盖用户疑问", "保留智能体问答入口，让系统既能“识别”，也能“解释”"],
        ],
        [4.4, 11.6],
    )
    add_picture(document, EXPERIMENT_DIR / "3e9e6b3fdc438e11d4fdf85042322c26.png", 15.5, "1 系统栏目化首页与个性推荐页面")

    add_heading(document, "二、算法原理介绍", level=1)
    add_heading(document, "（一）数据集准备与类别设计", level=2)
    add_body_paragraph(
        document,
        "本项目采用的是四类垃圾目标检测数据集，类别分别为 recyclable waste、hazardous waste、kitchen waste 和 other waste。数据集最终被整理为 YOLO 标准格式，并通过 dataset.yaml、train.txt、val.txt、test.txt 进行组织。相比只做图像分类，目标检测的优势在于能定位垃圾目标区域，更适合复杂背景下的网页识别演示。"
    )
    add_body_paragraph(
        document,
        "在数据准备阶段，我们首先统一了类别体系，再将训练、验证、测试三部分列表写入文本文件，便于 Linux 服务器直接读取。这样做的好处是训练流程清晰、可复现，也方便后续更换模型或重新训练。"
    )
    add_triple_line_table(
        document,
        "3 四类垃圾标签映射关系",
        ["英文标签", "中文解释", "说明"],
        [
            ["recyclable waste", "可回收垃圾", "主要包括纸箱、塑料盒、塑料瓶、金属罐、玻璃瓶等可资源化利用物品"],
            ["hazardous waste", "有害垃圾", "主要包括废电池、过期药品、废灯管等需要专门回收处理的物品"],
            ["kitchen waste", "厨余垃圾", "主要包括果皮菜叶、剩饭剩菜等易腐有机垃圾"],
            ["other waste", "其他垃圾", "不属于前三类且当前阶段较难资源化利用的剩余垃圾"],
        ],
        [4.2, 3.4, 8.4],
    )
    add_triple_line_table(
        document,
        "4 数据集划分结果",
        ["数据子集", "样本数量", "说明"],
        [
            ["训练集", str(info["train_count"]), "用于 YOLOv8 参数学习与主要特征提取"],
            ["验证集", str(info["val_count"]), "用于训练过程中的性能监控与早停判断"],
            ["测试集", str(info["test_count"]), "用于保留独立评估与后期案例核验"],
            ["总计", str(info["total_count"]), "四类垃圾检测任务总样本数"],
        ],
        [3.2, 3.0, 9.8],
    )
    add_picture(document, EXPERIMENT_DIR / "0d71f6c9b51da10ad560ee49a6fbdb33.png", 14.8, "2 数据集页面与类别说明")
    add_picture(document, EXPERIMENT_DIR / "train_batch0.jpg", 15.5, "3 训练样本与标注框可视化")

    add_heading(document, "（二）YOLOv8 检测原理", level=2)
    add_body_paragraph(
        document,
        "本项目的核心识别模型采用 YOLOv8s。该模型属于单阶段目标检测算法，能够在较快速度下同时完成目标定位与类别判断。相较于传统两阶段检测方法，YOLOv8 更适合课程项目中的实时网页展示，因为它在推理时延、工程集成和训练便利性方面更有优势。"
    )
    add_body_paragraph(
        document,
        "从结构上看，YOLOv8 可以简要分为三个部分：第一部分是 Backbone，用于从输入图像中提取不同尺度的特征；第二部分是 Neck，用于对多尺度特征进行融合，提升对大小目标的检测能力；第三部分是 Detection Head，用于输出类别概率、边界框和回归相关结果。本项目在训练中使用了官方预训练权重 yolov8s.pt，因此属于迁移学习方案，而不是从零开始训练。"
    )
    add_triple_line_table(
        document,
        "5 YOLOv8 关键模块与作用",
        ["模块", "主要作用", "在本项目中的意义"],
        [
            ["Backbone", "提取图像的深层语义特征", "帮助模型区分纸箱、果皮、塑料制品、电池等外观差异"],
            ["Neck", "融合多尺度特征信息", "提升复杂背景和不同尺寸垃圾目标的检测稳定性"],
            ["Detection Head", "输出类别和边界框结果", "直接生成网页展示所需的识别标签、置信度和目标框"],
            ["迁移学习", "使用预训练模型初始化参数", "减少训练时间，提高课程项目可落地性"],
        ],
        [2.8, 4.6, 8.6],
    )

    add_heading(document, "（三）智能体解释链路", level=2)
    add_body_paragraph(
        document,
        "仅有检测结果还不够。很多用户真正关心的问题是：“为什么它属于这一类？”“是不是直接丢进对应垃圾桶就可以？”“像充电宝、插头这类物品需要怎么处理？”为此，本项目在识别页面下方接入了阿里云通义千问模型，对识别结果进行二次解释。"
    )
    add_body_paragraph(
        document,
        "具体做法是：先由 YOLOv8 返回综合类别、目标明细和置信度，再把这些结构化信息连同用户问题一起发送给千问模型。智能体会结合垃圾分类常识与投放规则，输出更自然、更完整的解释文本。也就是说，本项目中的大模型不是替代视觉模型，而是承担“解释器”和“答疑助手”的角色。"
    )
    add_triple_line_table(
        document,
        "6 智能体输入输出设计",
        ["输入项", "内容", "输出结果"],
        [
            ["视觉结果", "综合类别、目标明细、置信度", "解释该垃圾为什么属于某一类"],
            ["用户问题", "例如是否能直接投放、是否需要送到专门回收点", "给出更专业的投放建议和处理说明"],
            ["调用模型", "阿里云通义千问 API（兼容 chat/completions 接口）", "返回自然语言分析结果，增强系统可解释性"],
        ],
        [3.0, 6.0, 7.0],
    )
    add_picture(document, EXPERIMENT_DIR / "340ea643a858e766e8273c70bdb056bd.png", 15.2, "4 垃圾分类智能助手界面")

    add_heading(document, "（四）项目涉及的关键公式", level=2)
    add_body_paragraph(
        document,
        "为了让项目说明书不只是描述功能，本节把训练、检测、推荐、账号安全和博客主题过滤中实际用到或直接相关的公式统一列出。这样在汇报时可以更清楚地说明：网页上的结果并不是简单写死的文字，而是由模型置信度、目标框重合度、用户偏好和互动行为共同计算得到的。"
    )
    add_body_paragraph(
        document,
        "目标框去重和结果稳定性判断中使用了 IoU 指标，公式为：IoU(A,B)=area(A∩B)/area(A∪B)。当两个同名检测框 IoU 大于 0.6 时，系统会保留置信度更高的结果，减少同一垃圾被重复标注的问题。若最高置信度低于 0.45，或多个候选框置信度非常接近，系统会给出“待确认”，提示用户重新拍摄或人工复核。"
    )
    add_body_paragraph(
        document,
        "个性推荐部分采用规则加权评分。对视频 v，推荐分数可写为：Score(v)=42+18I(hT>0)+4hT+16I(hR>0)+3hR+8I(hC>0)+10I(liked)+12I(favorite)-4I(viewed)+2I(featured)。其中 hT 表示视频主题与用户偏好主题的命中数量，hR 表示视频主题与近期识别历史映射主题的命中数量，hC 表示视频关注类别与近期识别类别是否命中。这样可以让用户最近识别过的垃圾类型影响后续视频推荐。"
    )
    add_body_paragraph(
        document,
        "账号系统没有直接保存明文密码，而是使用 password_hash=PBKDF2-HMAC-SHA256(password,salt,200000) 生成密码摘要。点赞、收藏和评论数则按用户行为统计，例如 LikeCount(v)=sum_u I(v∈LikedVideos(u))，CommentCount(v)=sum_c I(c.video_id=v)。博客内容过滤使用 KeywordHit(text)=I(存在 k∈K 且 k 出现在 text 中)，保证博客栏目和评论都围绕垃圾分类主题。"
    )
    add_triple_line_table(
        document,
        "7 项目关键公式与作用",
        ["公式/规则", "作用", "在系统中的位置"],
        [
            ["IoU(A,B)=area(A∩B)/area(A∪B)", "衡量两个检测框的重合程度", "检测框去重、结果稳定性判断"],
            ["Precision=TP/(TP+FP)", "衡量预测为正的目标中有多少是真的", "YOLOv8 训练性能评估"],
            ["Recall=TP/(TP+FN)", "衡量真实目标中有多少被模型检出", "YOLOv8 训练性能评估"],
            ["APc=∫P(R)dR，mAP=(1/C)sum(APc)", "汇总不同类别的检测平均精度", "YOLOv8 mAP50 与 mAP50-95 指标"],
            ["Score(v)=42+18I(hT>0)+4hT+16I(hR>0)+3hR+8I(hC>0)+10I(liked)+12I(favorite)-4I(viewed)+2I(featured)", "根据偏好、历史识别和互动行为计算视频推荐分", "个性推荐栏目"],
            ["PBKDF2-HMAC-SHA256(password,salt,200000)", "把用户密码转为不可逆摘要", "登录注册与账号安全"],
            ["KeywordHit(text)=I(存在 k∈K 且 k 出现在 text 中)", "限制博客和评论必须围绕垃圾分类主题", "分类博客栏目"],
        ],
        [5.5, 5.3, 5.2],
    )

    add_heading(document, "三、算法编写与调试说明", level=1)
    add_heading(document, "（一）YOLOv8 训练过程", level=2)
    add_body_paragraph(
        document,
        "模型训练在 Linux 服务器上完成。训练时使用官方预训练权重 yolov8s.pt 作为初始化参数，图像尺寸设置为 640，batch size 设置为 8，workers 设置为 4，同时开启 AMP 混合精度训练以提高训练效率。为了减少无效训练时间，我们将最大 epoch 设为 1200，并启用 patience=30 的早停机制，结果模型在第 132 轮自动提前停止。"
    )
    add_body_paragraph(
        document,
        "从实际结果看，虽然参数里写的是 1200 epoch，但这并不意味着模型一定会训满。早停的优势就在于：当验证集性能在连续一段时间内不再提升时，系统可以自动停止训练，避免时间浪费和过拟合。在本次实验中，mAP50 的最佳点出现在第 82 轮，mAP50-95 的最佳点出现在第 102 轮，说明较优性能实际上在中前期就已经出现。"
    )
    add_triple_line_table(
        document,
        "8 YOLOv8 训练参数设置",
        ["参数项", "取值", "说明"],
        [
            ["基础模型", "yolov8s.pt", "使用官方预训练权重进行迁移学习"],
            ["最大 epoch", str(info["args"]["epochs"]), "配合早停机制，实际训练在第 132 轮停止"],
            ["patience", str(info["args"]["patience"]), "连续 30 轮无明显提升时提前停止"],
            ["batch size", str(info["args"]["batch"]), "平衡显存占用与训练速度"],
            ["imgsz", str(info["args"]["imgsz"]), "训练输入分辨率为 640"],
            ["device", str(info["args"]["device"]), "使用 Linux 服务器 GPU 训练"],
            ["optimizer", str(info["args"]["optimizer"]), "由 Ultralytics 自动选择优化器"],
            ["AMP", "开启", "提高训练效率并减少显存占用"],
        ],
        [3.3, 3.0, 9.7],
    )
    curve_path = create_training_curve(EXPERIMENT_DIR / "results.csv", TRAINING_CURVE)
    if curve_path:
        add_picture(document, curve_path, 15.5, "5 训练过程中 Precision、Recall 与 mAP 指标变化曲线")
    add_picture(document, EXPERIMENT_DIR / "BoxPR_curve.png", 14.6, "6 YOLOv8 检测结果 PR 曲线")
    add_picture(document, EXPERIMENT_DIR / "val_batch0_pred.jpg", 15.5, "7 验证集预测结果示例")

    add_heading(document, "（二）系统实现过程与技术重点", level=2)
    add_body_paragraph(
        document,
        "在工程实现层面，本项目并没有只做一个“识别按钮 + 输出结果”的单页 Demo，而是进一步搭建了一个栏目化网页系统。用户需要先登录注册，然后可以在垃圾识别、个性推荐、分类博客和个人信息四个栏目之间切换。这样的设计更接近真实产品，也更利于课程汇报时展示完整项目流程。"
    )
    add_body_paragraph(
        document,
        "网页前端采用了比较明显的卡片式布局，每个栏目都用独立信息块组织内容，视觉层次清楚，适合课堂演示。这部分页面结构和样式是在 Codex 协助下完成的，可以看出整体风格比较统一：绿色渐变导航、浅色背景、卡片化内容区、独立的评论与互动模块。对课程项目来说，这样的界面比传统的“纯表单 + 纯文字”更有展示效果。"
    )
    add_triple_line_table(
        document,
        "9 系统模块与实现重点",
        ["模块", "功能", "技术重点"],
        [
            ["垃圾识别栏目", "上传图片、返回检测框和四类垃圾结果", "对接自训练 YOLOv8 权重，并输出中文解释"],
            ["智能体助手", "对识别结果进行问答分析", "调用阿里云通义千问 API，补充特殊物品投放规则"],
            ["个性推荐栏目", "展示垃圾分类相关视频", "使用真实视频链接和匹配封面，支持点赞、收藏和评论"],
            ["分类博客栏目", "发布垃圾分类实践内容", "支持发帖、点赞、评论，且主题限定为垃圾分类相关内容"],
            ["个人信息栏目", "展示头像、地区和个人互动记录", "汇总用户点赞、评论过的视频以及发布过的博客"],
        ],
        [3.2, 5.0, 7.8],
    )
    add_picture(document, EXPERIMENT_DIR / "221ab6a4942f16226c7eedd835fc2e6d.png", 15.2, "8 个性推荐栏目界面")
    add_picture(document, EXPERIMENT_DIR / "afa6b95d4965b3102c240e921107015f.png", 15.2, "9 分类博客栏目界面")
    add_picture(document, EXPERIMENT_DIR / "a4cf9858d7a1fec6ede573ccb0a7be9f.png", 15.2, "10 个人信息栏目界面")

    add_heading(document, "（三）调试过程中的关键问题", level=2)
    add_body_paragraph(
        document,
        "项目制作过程中最值得总结的并不只是“模型训出来了”，而是训练与调试过程中对垃圾分类知识本身也有了更深的理解。一个很有意思的例子是，在进行性能检测时，我曾经放入一张插头图片，模型给出的结果是“可回收垃圾”。当时第一反应是模型出问题了，因为我原本以为这类电子废品应该属于有害垃圾或者不可回收垃圾。"
    )
    add_body_paragraph(
        document,
        "后来查阅资料并结合当前数据集的分类标准后发现，这类物品并不一定要简单地归入“有害垃圾”。像插头这类以金属和塑料为主、没有明显有毒污染成分的器件，在很多实际分类规则中更接近可回收或电子回收处理对象。也就是说，这次误判怀疑反而促使我重新学习了垃圾分类知识，这一点很有意思，也很符合课程大作业“在做项目中学习”的要求。"
    )
    add_body_paragraph(
        document,
        "另一个典型例子是充电宝。对于这种特殊可回收垃圾，仅靠检测标签还不够，因为用户真正想知道的是“能不能直接扔进普通可回收桶”。在系统中，用户可以继续追问智能体助手，智能体会给出更加专业的解释，例如充电宝虽然可以归入可回收体系，但通常不建议直接投入普通可回收垃圾桶，而应尽量送至社区或学校设置的电子废弃物回收点。这也体现了视觉模型与大模型协同的意义：前者负责识别，后者负责解释。"
    )
    add_body_paragraph(
        document,
        "在工程层面，我们还遇到过网页链路与模型直推结果不一致的问题。继续排查后发现，根因并不是模型本身，而是图像通道顺序在某一段处理链路中出现了 RGB/BGR 混用。修复这一问题后，网页端结果和模型直接推理结果重新保持一致，说明完整系统调试中“数据流一致性”同样是一个重要技术点。"
    )

    add_heading(document, "四、结果展示", level=1)
    add_heading(document, "（一）训练性能结果", level=2)
    add_body_paragraph(
        document,
        "从实验结果看，模型在 132 轮训练后停止，整体性能达到较稳定水平。对于课程大作业而言，这样的效果已经可以支撑网页演示和案例展示。尤其是 mAP50 达到 0.90256，说明模型在较宽松评价标准下具备较好的检测能力；mAP50-95 达到 0.63098，也说明其在更严格 IoU 条件下仍保留了一定稳定性。"
    )
    add_triple_line_table(
        document,
        "10 训练核心指标统计",
        ["指标项", "最优值", "对应 epoch", "说明"],
        [
            ["Precision", f"{float(info['best_precision']['metrics/precision(B)']):.5f}", info["best_precision"]["epoch"], "精度最高出现在训练后期，误检控制较好"],
            ["Recall", f"{float(info['best_recall']['metrics/recall(B)']):.5f}", info["best_recall"]["epoch"], "召回率最高时模型能检出更多目标"],
            ["mAP50", f"{float(info['best_map50']['metrics/mAP50(B)']):.5f}", info["best_map50"]["epoch"], "宽松 IoU 标准下的综合检测性能"],
            ["mAP50-95", f"{float(info['best_map95']['metrics/mAP50-95(B)']):.5f}", info["best_map95"]["epoch"], "严格标准下的综合性能，更能反映泛化能力"],
            ["最终 mAP50", f"{float(info['final_row']['metrics/mAP50(B)']):.5f}", info["final_row"]["epoch"], "训练停止时的最终结果"],
            ["最终 mAP50-95", f"{float(info['final_row']['metrics/mAP50-95(B)']):.5f}", info["final_row"]["epoch"], "训练停止时的严格评估结果"],
        ],
        [3.0, 3.0, 2.5, 8.5],
    )

    add_heading(document, "（二）识别案例展示", level=2)
    add_body_paragraph(
        document,
        "在实际网页展示中，系统能够对常见垃圾图片给出较清晰的识别结果，并同步展示综合判断、检测明细和解释建议。下图分别展示了塑料瓶和果皮类垃圾的识别效果。这样的展示方式比单纯打印类别标签更直观，也更适合课堂汇报。"
    )
    add_picture(document, EXPERIMENT_DIR / "cd90216560a5b9a36b71eac4ab3a869d.png", 15.2, "11 塑料瓶识别结果展示")
    add_picture(document, EXPERIMENT_DIR / "dad89d506fa110ecbe5c5706bbc08116.png", 15.2, "12 果皮类垃圾识别结果展示")

    add_heading(document, "（三）系统功能验证结果", level=2)
    add_body_paragraph(
        document,
        "除了模型识别精度外，本项目还重点验证了系统功能完整性，包括用户登录、识别、智能体解释、视频推荐、博客互动和个人信息管理等。对于课程展示来说，完整的系统流程往往比单一算法结果更能体现项目完成度。"
    )
    add_triple_line_table(
        document,
        "11 系统功能验证情况",
        ["功能点", "验证结果", "说明"],
        [
            ["账号登录注册", "通过", "支持保存用户账号、密码摘要、头像与地区信息"],
            ["垃圾图片识别", "通过", "网页端可正确加载自训练 best.pt 并返回四类垃圾结果"],
            ["智能体问答", "通过", "接入阿里云通义千问，能对特殊垃圾给出补充解释"],
            ["个性推荐", "通过", "推荐栏目的视频均与垃圾分类相关，并带有真实封面"],
            ["博客栏目", "通过", "支持垃圾分类主题发帖、点赞、评论和个人历史查看"],
            ["个人信息页", "通过", "可查看自己点赞、评论过的视频和发布过的博客"],
        ],
        [3.2, 2.2, 11.6],
    )

    add_heading(document, "五、GitHub 与 Render 公网部署", level=1)
    add_heading(document, "（一）上传到 GitHub 的准备", level=2)
    add_body_paragraph(
        document,
        "为了让项目可以被 Render 拉取和自动部署，需要先把本地项目上传到 GitHub。根据 GitHub 官方文档，已有本地项目上传前应先在 GitHub 创建一个空仓库，再在本地执行 git init、git add、git commit、git remote add 和 git push 等命令。特别需要注意的是，官方文档也提醒不要把密码、API Key 等敏感信息提交到远程仓库，因此本项目新增了 .gitignore，将 .env、uploads、运行数据库和临时文件排除在版本控制之外。"
    )
    add_body_paragraph(
        document,
        "本项目需要提交的核心文件包括 app 目录、requirements.txt、README.md、render.yaml、.python-version、.env.example 以及自训练权重 artifacts/weights/garbage4_best.pt。该权重约 22MB，未超过 GitHub 单文件 100MB 限制，适合作为课程演示直接随仓库提交；如果后续模型超过限制，则可以改用 Git LFS、GitHub Release 或对象存储下载地址。"
    )
    add_triple_line_table(
        document,
        "12 GitHub 上传步骤",
        ["步骤", "命令/操作", "说明"],
        [
            ["1", "在 GitHub 新建空仓库", "不要初始化 README、License 或 .gitignore，避免首次 push 冲突"],
            ["2", "cd /d E:\\人工智能\\garbage_ai_system", "进入项目根目录"],
            ["3", "git init", "初始化本地 Git 仓库"],
            ["4", "git add . && git commit -m \"init garbage classification web app\"", "提交代码、前端、后端、部署配置和模型权重"],
            ["5", "git branch -M main", "统一主分支名称"],
            ["6", "git remote add origin https://github.com/用户名/garbage-ai-system.git", "绑定远程仓库地址"],
            ["7", "git push -u origin main", "推送到 GitHub，后续提交可直接 git push"],
        ],
        [1.8, 7.0, 7.2],
    )

    add_heading(document, "（二）Render 公网部署方案", level=2)
    add_body_paragraph(
        document,
        "Render 官方文档说明，Web Service 会为项目分配公开的 onrender.com 子域名，并且服务必须绑定 0.0.0.0 上的端口；实际端口建议读取 PORT 环境变量。本项目已经加入 render.yaml，启动命令为 python -m uvicorn app.main:app --host 0.0.0.0 --port $PORT，并提供 /health 健康检查接口。Render 连接 GitHub 仓库后，每次 main 分支更新都可以自动重新构建和部署。"
    )
    add_body_paragraph(
        document,
        "由于 Render 新建 Python 服务的默认版本可能变化，本项目还新增了 .python-version，并固定为 3.11.11，避免 torch、ultralytics 等依赖在过新的 Python 版本上出现兼容问题。部署时不要上传本地 .env，而是在 Render 控制台 Environment 页面中填写阿里云 API Key 等环境变量。"
    )
    add_triple_line_table(
        document,
        "13 Render 部署参数与环境变量",
        ["配置项", "推荐取值", "说明"],
        [
            ["Runtime", "Python 3", "使用 Render 原生 Python 运行时"],
            ["Build Command", "pip install -r requirements.txt", "安装 FastAPI、Ultralytics、Pillow 等依赖"],
            ["Start Command", "python -m uvicorn app.main:app --host 0.0.0.0 --port $PORT", "绑定 Render 提供的公网端口"],
            ["Health Check Path", "/health", "用于检查服务是否启动成功"],
            ["LLM_API_BASE", "https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions", "阿里云通义千问兼容接口地址"],
            ["LLM_MODEL", "qwen-plus", "智能体问答所用模型，可按账号权限调整"],
            ["LLM_API_KEY", "在 Render 环境变量中手动填写", "不能写入代码仓库或说明书截图"],
            ["VISION_MODEL_NAME", "garbage4_best.pt", "默认从 artifacts/weights 中加载自训练权重"],
        ],
        [3.5, 6.5, 6.0],
    )
    add_body_paragraph(
        document,
        "部署完成后，Render 会生成一个形如 https://项目名.onrender.com 的公网地址。测试时应依次检查首页是否能打开、/health 是否返回 ok、图片识别是否能加载 garbage4_best.pt、智能体是否能正常调用阿里云 API。需要说明的是，Render 免费服务的文件系统默认是临时的，重新部署或重启后运行时生成的 SQLite 账号数据、评论和上传图片可能丢失。课程演示阶段可以接受这一点；如果后续要长期运行，则应改用 Render Disk 或外部数据库保存用户数据。"
    )

    add_heading(document, "六、成员信息", level=1)
    add_body_paragraph(
        document,
        "由于最终小组成员名单和学号信息需要以实际提交版本为准，下面先保留了待补充表格。正式上交前只需把相应信息替换即可，不影响正文内容结构。"
    )
    add_triple_line_table(
        document,
        "14 小组成员信息表",
        ["序号", "姓名", "学号", "主要分工"],
        [
            ["1", "待补充", "待补充", "项目统筹、模型训练"],
            ["2", "待补充", "待补充", "网页前端与交互设计"],
            ["3", "待补充", "待补充", "数据整理、测试与结果分析"],
            ["4", "待补充", "待补充", "PPT 汇报与文档整理"],
        ],
        [2.0, 3.5, 4.0, 7.0],
    )

    add_heading(document, "七、总结", level=1)
    add_body_paragraph(
        document,
        "总体来看，本项目不仅完成了基于 YOLOv8 的校园垃圾识别系统，还进一步结合了阿里云通义千问智能体、个性推荐、博客互动和个人信息管理等模块，形成了一个完整的垃圾分类学习与应用平台。项目的技术重点主要体现在四个方面：第一，使用迁移学习高效训练 YOLOv8；第二，将数据集整理为标准 YOLO 检测格式；第三，用大模型补足“解释”和“答疑”能力；第四，把算法结果嵌入到一个更适合展示的栏目化网页中。"
    )
    add_body_paragraph(
        document,
        "对我来说，这个项目最有价值的地方在于，它并不是单纯“把模型跑出来”，而是在不断训练、调试、查资料和修改网页的过程中，对垃圾分类本身也建立了更加扎实的理解。尤其是插头、充电宝这类边界物品带来的疑问，让我意识到人工智能项目不仅是在解决技术问题，也是在推动我们重新理解现实世界中的规则和知识。"
    )

    document.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    output = build_report()
    print(output)
